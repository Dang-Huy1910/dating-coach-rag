from __future__ import annotations

import csv
import html
import io
import re
from pathlib import Path

ALLOWED_EXTENSIONS = {
    ".md",
    ".markdown",
    ".txt",
    ".pdf",
    ".docx",
    ".html",
    ".htm",
    ".csv",
}

FORMAT_LABELS = [
    {"ext": ".md", "label": "Markdown", "note": "Hướng dẫn / ghi chú có tiêu đề"},
    {"ext": ".txt", "label": "Plain text", "note": "Văn bản thuần"},
    {"ext": ".pdf", "label": "PDF", "note": "PDF có lớp chữ (không OCR ảnh scan)"},
    {"ext": ".docx", "label": "Word", "note": "Microsoft Word .docx"},
    {"ext": ".html", "label": "HTML", "note": "Trang HTML / bài viết đã lưu"},
    {"ext": ".csv", "label": "CSV", "note": "Bảng ghi chú dạng CSV"},
]

_TAG = re.compile(r"<[^>]+>")
_WS = re.compile(r"\n{3,}")


def is_allowed_filename(name: str) -> bool:
    return Path(name).suffix.lower() in ALLOWED_EXTENSIONS


def extract_text(filename: str, data: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix not in ALLOWED_EXTENSIONS:
        raise ValueError(f"Định dạng không hỗ trợ: {suffix or '(không có đuôi)'}")
    if suffix in {".md", ".markdown", ".txt"}:
        return _decode_text(data)
    if suffix == ".pdf":
        return _extract_pdf(data)
    if suffix == ".docx":
        return _extract_docx(data)
    if suffix in {".html", ".htm"}:
        return _extract_html(data)
    if suffix == ".csv":
        return _extract_csv(data)
    raise ValueError(f"Định dạng không hỗ trợ: {suffix}")


def _decode_text(data: bytes) -> str:
    for enc in ("utf-8", "utf-8-sig", "cp1258", "latin-1"):
        try:
            text = data.decode(enc)
            break
        except UnicodeDecodeError:
            continue
    else:
        text = data.decode("utf-8", errors="replace")
    text = text.strip()
    if not text:
        raise ValueError("File trống hoặc không đọc được chữ.")
    return text


def _extract_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ValueError("Thiếu thư viện pypdf để đọc PDF.") from exc
    reader = PdfReader(io.BytesIO(data))
    parts: list[str] = []
    for page in reader.pages:
        piece = (page.extract_text() or "").strip()
        if piece:
            parts.append(piece)
    text = "\n\n".join(parts).strip()
    if not text:
        raise ValueError("PDF không có lớp chữ (có thể là ảnh scan — chưa hỗ trợ OCR).")
    return text


def _extract_docx(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise ValueError("Thiếu thư viện python-docx để đọc Word.") from exc
    doc = Document(io.BytesIO(data))
    parts = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n\n".join(parts).strip()
    if not text:
        raise ValueError("File Word trống hoặc không có đoạn văn bản.")
    return text


def _extract_html(data: bytes) -> str:
    raw = _decode_text(data)
    # Drop script/style blocks then strip tags.
    raw = re.sub(r"(?is)<(script|style).*?>.*?</\1>", " ", raw)
    text = _TAG.sub(" ", raw)
    text = html.unescape(text)
    text = re.sub(r"[ \t]+", " ", text)
    text = _WS.sub("\n\n", text).strip()
    if not text:
        raise ValueError("HTML không có nội dung chữ.")
    return text


def _extract_csv(data: bytes) -> str:
    raw = _decode_text(data)
    reader = csv.reader(io.StringIO(raw))
    lines: list[str] = []
    for row in reader:
        cells = [c.strip() for c in row if c and c.strip()]
        if cells:
            lines.append(" | ".join(cells))
    text = "\n".join(lines).strip()
    if not text:
        raise ValueError("CSV trống.")
    return text
